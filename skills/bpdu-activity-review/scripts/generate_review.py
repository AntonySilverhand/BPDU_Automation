#!/usr/bin/env python3
"""
BP Debate Union — Activity Review Document Generator

Generates a Word document (.docx) for weekly club activity reviews.
Embeds user-provided photos and writes a third-person description.

Usage:
    python generate_review.py --week 7 --date "11月12日" --location "博闻楼B-606" \
        --activity "苏格拉底式研讨会" --participants 20 \
        --topic "社交媒体的continuing the spark现象" \
        --description "11月12日，BP Debate Union在博闻楼B-606开展了..." \
        --photos "photo1.jpg" "photo2.jpg" "photo3.jpg"

All arguments except --description are editable constants at the top of INPUT DATA.
"""

import argparse
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===================== INPUT DATA =====================
# Edit these values, or pass via command-line arguments (command-line takes priority)
WEEK_NUMBER = 7          # 第几周
DATE = "11月12日"          # Date of the activity (last week's format, e.g. 上周三)
LOCATION = "博闻楼B-606"   # Venue
ACTIVITY_TYPE = "苏格拉底式研讨会"   # Type of activity
PARTICIPANTS = 20          # Number of participants
TOPIC = '社交媒体的"continuing the spark"现象'   # Discussion topic/motion
DESCRIPTION = (
    "上周三，BP Debate Union在博闻楼B-606开展了苏格拉底式研讨会。"
    "本次活动共有20名同学参加，围绕社交媒体延续火花现象展开讨论。"
    "同学们积极发言，热烈探讨，提升了批判性思维和表达能力。"
    "大家表示收获颇丰，期待下次参与。"
)
# Photo paths: provide 2-3 paths (minimum 2 required)
# Photos are embedded in the document in the order provided.
PHOTO_PATHS = []          # e.g. ["photo1.jpg", "photo2.jpg", "photo3.jpg"]
# =====================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(SKILL_DIR, "output")


def parse_args():
    parser = argparse.ArgumentParser(description="Generate BP DU Activity Review Word document.")
    parser.add_argument("--week", type=int, default=WEEK_NUMBER)
    parser.add_argument("--date", type=str, default=DATE)
    parser.add_argument("--location", type=str, default=LOCATION)
    parser.add_argument("--activity", type=str, default=ACTIVITY_TYPE)
    parser.add_argument("--participants", type=int, default=PARTICIPANTS)
    parser.add_argument("--topic", type=str, default=TOPIC)
    parser.add_argument("--description", type=str, default=DESCRIPTION)
    parser.add_argument("--photos", nargs="*", default=PHOTO_PATHS,
                        help="Paths to 2-3 photos (minimum 2 required)")
    return parser.parse_args()


def generate_filename(week_number):
    return f"国际学院BP_Debate_Union_第{week_number}周活动剪影.docx"


def make_unique_path(filepath):
    """If file exists, append timestamp to avoid overwrite."""
    if not os.path.exists(filepath):
        return filepath
    base, ext = os.path.splitext(filepath)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{base}_{timestamp}{ext}"


def generate(args):
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Validate photo count
    valid_photos = [p for p in args.photos if os.path.exists(p)]
    if len(valid_photos) < 2:
        print(f"WARNING: Only {len(valid_photos)} valid photo(s) found. Minimum 2 required.", file=sys.stderr)
        if len(valid_photos) == 0:
            print("ERROR: No valid photos. Cannot generate document.", file=sys.stderr)
            return None

    filename = generate_filename(args.week)
    filepath = make_unique_path(os.path.join(OUTPUT_DIR, filename))

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"国际学院BP Debate Union第{args.week}周活动剪影")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Photos
    for photo in valid_photos:
        doc.add_picture(photo, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # spacing

    # Description
    desc_para = doc.add_paragraph(args.description)
    desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(filepath)
    print(f"Generated: {filepath}")
    return filepath


if __name__ == "__main__":
    args = parse_args()
    result = generate(args)
    if result is None:
        sys.exit(1)
